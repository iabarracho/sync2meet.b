from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.text.paragraph import Paragraph
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas


_INLINE_MARKDOWN = re.compile(
    r"(\*\*(.+?)\*\*|\*(.+?)\*|_(.+?)_)"
)
_BULLET_RE = re.compile(r"^(?:[-*•–—]\s+)(.+)$")
_NUMBERED_RE = re.compile(r"^(\d+)[.)]\s+(.+)$")


def _parse_table_cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _is_table_separator(cells: list[str]) -> bool:
    if not cells:
        return True
    return all(re.fullmatch(r":?-{2,}:?", cell.replace(" ", "")) for cell in cells if cell)


def _add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    rows: list[list[str]] = []
    for line in table_lines:
        cells = _parse_table_cells(line)
        if _is_table_separator(cells):
            continue
        rows.append(cells)
    if not rows:
        return

    col_count = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx in range(col_count):
            value = row[col_idx] if col_idx < len(row) else ""
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = ""
            paragraph = cell.paragraphs[0]
            _add_inline_runs(paragraph, value)


def _add_inline_runs(paragraph: Paragraph, text: str) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)

    pos = 0
    for match in _INLINE_MARKDOWN.finditer(text):
        if match.start() > pos:
            paragraph.add_run(text[pos : match.start()])
        if match.group(2) is not None:
            run = paragraph.add_run(match.group(2))
            run.bold = True
        elif match.group(3) is not None:
            run = paragraph.add_run(match.group(3))
            run.italic = True
        else:
            run = paragraph.add_run(match.group(4))
            run.italic = True
        pos = match.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])
    if not text:
        paragraph.add_run("")


def export_docx(content: str, output_path: Path) -> Path:
    doc = Document()
    lines = content.splitlines()
    index = 0
    while index < len(lines):
        raw = lines[index]
        line = raw.strip()
        if not line:
            doc.add_paragraph("")
            index += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            index += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            index += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            index += 1
            continue
        if line.startswith("|") and line.endswith("|"):
            table_lines: list[str] = []
            while index < len(lines):
                candidate = lines[index].strip()
                if not (candidate.startswith("|") and candidate.endswith("|")):
                    break
                table_lines.append(candidate)
                index += 1
            _add_markdown_table(doc, table_lines)
            continue

        bullet = _BULLET_RE.match(line)
        if bullet:
            para = doc.add_paragraph(style="List Bullet")
            _add_inline_runs(para, bullet.group(1))
            index += 1
            continue

        numbered = _NUMBERED_RE.match(line)
        if numbered:
            para = doc.add_paragraph(style="List Number")
            _add_inline_runs(para, numbered.group(2))
            index += 1
            continue

        para = doc.add_paragraph()
        _add_inline_runs(para, line)
        index += 1

    doc.save(str(output_path))
    return output_path


def export_pdf(content: str, output_path: Path) -> Path:
    c = canvas.Canvas(str(output_path), pagesize=A4)
    width, height = A4
    x, y = 40, height - 40
    line_height = 14
    c.setFont("Helvetica", 10)
    for raw in content.splitlines():
        line = raw.strip() or " "
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        line = re.sub(r"(?<!\*)\*(.+?)\*(?!\*)", r"\1", line)
        chunks = _wrap(line, 90)
        for chunk in chunks:
            if y < 40:
                c.showPage()
                c.setFont("Helvetica", 10)
                y = height - 40
            c.drawString(x, y, chunk)
            y -= line_height
    c.save()
    return output_path


def _wrap(text: str, width: int) -> list[str]:
    words = text.split()
    if not words:
        return [""]
    lines: list[str] = []
    current: list[str] = []
    for w in words:
        test = " ".join(current + [w])
        if len(test) <= width:
            current.append(w)
        else:
            if current:
                lines.append(" ".join(current))
            current = [w]
    if current:
        lines.append(" ".join(current))
    return lines


def render_agenda_variables(
    template: str,
    client_name: str,
    meeting_date: str,
    suggestions: str = "",
) -> str:
    content = template.replace("[NOME CLIENTE]", client_name)
    content = content.replace("[DATA]", meeting_date)
    if suggestions:
        for marker in ("- XX", "• XX"):
            if marker in content:
                content = content.replace(marker, suggestions, 1)
                break
    return content


def render_minutes_variables(
    template: str,
    client_name: str,
    meeting_date: str,
    participants: str,
) -> str:
    content = template.replace("[NOME CLIENTE]", client_name.upper())
    content = content.replace("[DATA]", meeting_date)
    content = content.replace("[PARTICIPANTES]", participants)
    return content


def short_summary_from_content(content: str, max_len: int = 280) -> str:
    plain = re.sub(r"[#*|_\[\]]", "", content)
    plain = " ".join(plain.split())
    if len(plain) <= max_len:
        return plain
    return plain[: max_len - 3] + "..."
